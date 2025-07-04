from django.db import models
from django.core.validators import MinValueValidator, MaxValueValidator
from django.utils.translation import gettext_lazy as _
from django.utils import timezone
from django.db.models import (
    Sum, F, ExpressionWrapper,
    DurationField, Q, Count, Case, When
)
from django.db.models.functions import (
    ExtractMonth, ExtractYear, ExtractQuarter
)
from django.core.exceptions import ValidationError
from django.urls import reverse
from django.conf import settings
from django.core.files.base import ContentFile
from django.template.loader import render_to_string
from contracts.models import Contract
from users.models import User
from datetime import timedelta, datetime
import pandas as pd
from docx import Document
from openpyxl import Workbook
from io import BytesIO
import json
import logging
from import_export import resources
from django_ckeditor_5.fields import CKEditor5Field
from django_tables2 import tables
from django_filters import FilterSet

logger = logging.getLogger(__name__)


class TimeStampedModel(models.Model):
    """Абстрактная модель с временными метками"""
    created_at = models.DateTimeField(auto_now_add=True, verbose_name=_("Дата создания"))
    updated_at = models.DateTimeField(auto_now=True, verbose_name=_("Дата обновления"))

    class Meta:
        abstract = True


class RentalPeriodQuerySet(models.QuerySet):
    """Кастомный QuerySet для аналитики аренды"""

    def active(self):
        return self.filter(
            end_date__gte=timezone.now().date(),
            contract__signed=True
        )

    def by_asset_type(self, asset_type):
        return self.filter(contract__asset__asset_type=asset_type)

    def with_profit(self):
        return self.annotate(
            profit=ExpressionWrapper(
                F('contract__price') * F('duration_days') / 30,
                output_field=models.DecimalField()
            )
        )

    def to_dataframe(self):
        """Конвертация в Pandas DataFrame"""
        values = self.annotate(
            month=ExtractMonth('start_date'),
            year=ExtractYear('start_date'),
            asset_type=F('contract__asset__asset_type'),
            client_name=F('contract__client__name')
        ).values(
            'id', 'start_date', 'end_date',
            'actual_price', 'month', 'year',
            'asset_type', 'client_name'
        )
        return pd.DataFrame.from_records(values)


class RentalPeriod(TimeStampedModel):
    """Модель периода аренды с расширенной аналитикой"""
    objects = RentalPeriodQuerySet.as_manager()

    class Meta:
        verbose_name = _("Период аренды")
        verbose_name_plural = _("Периоды аренды")
        ordering = ['-start_date']
        indexes = [
            models.Index(fields=['start_date', 'end_date']),
            models.Index(fields=['contract']),
        ]
        constraints = [
            models.CheckConstraint(
                check=Q(end_date__gte=F('start_date')),
                name='end_date_after_start_date'
            )
        ]

    contract = models.OneToOneField(
        Contract,
        on_delete=models.CASCADE,
        related_name='rental_period',
        verbose_name=_("Договор")
    )
    start_date = models.DateField(verbose_name=_("Дата начала"))
    end_date = models.DateField(verbose_name=_("Дата окончания"))
    actual_price = models.DecimalField(
        max_digits=12,
        decimal_places=2,
        verbose_name=_("Фактическая стоимость")
    )
    discount_applied = models.PositiveIntegerField(
        default=0,
        validators=[MaxValueValidator(100)],
        verbose_name=_("Скидка (%)")
    )
    notes = CKEditor5Field(
        verbose_name=_("Комментарии"),
        blank=True,
        config_name='extends'
    )

    @property
    def duration_days(self):
        """Длительность периода в днях"""
        return (self.end_date - self.start_date).days + 1

    @property
    def monthly_rate(self):
        """Средняя месячная стоимость"""
        return self.actual_price * 30 / self.duration_days

    def clean(self):
        """Валидация данных"""
        errors = {}
        if self.start_date > self.end_date:
            errors['end_date'] = _("Дата окончания не может быть раньше даты начала")

        if self.actual_price <= 0:
            errors['actual_price'] = _("Стоимость должна быть положительной")

        if errors:
            raise ValidationError(errors)

    def generate_docx_report(self):
        """Генерация отчета в формате Word"""
        doc = Document()
        doc.add_heading(f'Отчет по аренде #{self.id}', level=1)

        # Основная информация
        table = doc.add_table(rows=4, cols=2)
        table.cell(0, 0).text = 'Клиент'
        table.cell(0, 1).text = str(self.contract.client)
        table.cell(1, 0).text = 'Период'
        table.cell(1, 1).text = f"{self.start_date} - {self.end_date}"
        table.cell(2, 0).text = 'Стоимость'
        table.cell(2, 1).text = f"{self.actual_price:.2f} руб."
        table.cell(3, 0).text = 'Длительность'
        table.cell(3, 1).text = f"{self.duration_days} дней"

        # Сохраняем в BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def generate_excel_report(self):
        """Генерация отчета в Excel"""
        wb = Workbook()
        ws = wb.active
        ws.title = "Аренда"

        # Заголовки
        headers = [
            'ID', 'Клиент', 'Начало', 'Конец',
            'Дней', 'Стоимость', 'Ставка/мес'
        ]
        ws.append(headers)

        # Данные
        ws.append([
            self.id,
            str(self.contract.client),
            self.start_date,
            self.end_date,
            self.duration_days,
            float(self.actual_price),
            float(self.monthly_rate)
        ])

        buffer = BytesIO()
        wb.save(buffer)
        return buffer.getvalue()

    def save_report_to_file(self, format='docx'):
        """Сохранение отчета в файл"""
        if format == 'docx':
            content = self.generate_docx_report()
            filename = f'report_{self.id}.docx'
        elif format == 'xlsx':
            content = self.generate_excel_report()
            filename = f'report_{self.id}.xlsx'
        else:
            raise ValueError("Unsupported format")

        return ContentFile(content, name=filename)


class FinancialReport(TimeStampedModel):
    """Финансовый отчет с экспортом в разные форматы"""

    class Meta:
        verbose_name = _("Финансовый отчет")
        verbose_name_plural = _("Финансовые отчеты")
        ordering = ['-period_end']

    PERIOD_TYPES = [
        ('day', _('День')),
        ('week', _('Неделя')),
        ('month', _('Месяц')),
        ('quarter', _('Квартал')),
        ('year', _('Год'))
    ]

    title = models.CharField(max_length=255, verbose_name=_("Название"))
    period_type = models.CharField(
        max_length=10,
        choices=PERIOD_TYPES,
        default='month',
        verbose_name=_("Тип периода")
    )
    period_start = models.DateField(verbose_name=_("Начало периода"))
    period_end = models.DateField(verbose_name=_("Конец периода"))
    report_data = models.JSONField(
        verbose_name=_("Данные отчета"),
        default=dict
    )
    excel_file = models.FileField(
        upload_to='reports/excel/',
        null=True,
        blank=True,
        verbose_name=_("Excel файл")
    )
    word_file = models.FileField(
        upload_to='reports/word/',
        null=True,
        blank=True,
        verbose_name=_("Word файл")
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        verbose_name=_("Автор отчета")
    )

    def generate_excel(self):
        """Генерация Excel отчета с помощью pandas"""
        try:
            df = pd.DataFrame(self.report_data['data'])
            output = BytesIO()

            with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                df.to_excel(writer, sheet_name='Данные', index=False)

                # Добавляем графики
                workbook = writer.book
                worksheet = writer.sheets['Данные']

                chart = workbook.add_chart({'type': 'column'})
                chart.add_series({
                    'values': '=Данные!$B$2:$B$10',
                    'categories': '=Данные!$A$2:$A$10',
                })
                worksheet.insert_chart('D2', chart)

            self.excel_file.save(
                f'report_{self.id}.xlsx',
                ContentFile(output.getvalue())
            )
            return True
        except Exception as e:
            logger.error(f"Error generating Excel report: {e}")
            return False

    def generate_word(self):
        """Генерация Word отчета"""
        try:
            context = {
                'report': self,
                'data': self.report_data.get('data', [])
            }
            html = render_to_string('analytics/report_template.html', context)

            doc = Document()
            doc.add_heading(self.title, level=1)

            # Добавляем таблицу
            table = doc.add_table(rows=1, cols=len(self.report_data['headers']))
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(self.report_data['headers']):
                hdr_cells[i].text = str(header)

            for row in self.report_data['data']:
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)

            buffer = BytesIO()
            doc.save(buffer)

            self.word_file.save(
                f'report_{self.id}.docx',
                ContentFile(buffer.getvalue())
            )
            return True
        except Exception as e:
            logger.error(f"Error generating Word report: {e}")
            return False


class Dashboard(TimeStampedModel):
    """Интерактивный дашборд для визуализации данных"""

    class Meta:
        verbose_name = _("Дашборд")
        verbose_name_plural = _("Дашборды")
        ordering = ['title']

    title = models.CharField(max_length=255, verbose_name=_("Название"))
    description = CKEditor5Field(
        verbose_name=_("Описание"),
        blank=True,
        config_name='extends'
    )
    config = models.JSONField(
        verbose_name=_("Конфигурация"),
        default=dict,
        help_text=_("JSON конфигурация для дашборда")
    )
    is_default = models.BooleanField(
        default=False,
        verbose_name=_("Дашборд по умолчанию")
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        verbose_name=_("Автор")
    )

    def __str__(self):
        return self.title

    def get_plotly_config(self):
        """Получение конфигурации для Plotly"""
        return {
            'data': self.config.get('charts', []),
            'layout': self.config.get('layout', {})
        }


class RentalPeriodResource(resources.ModelResource):
    """Ресурс для импорта/экспорта периодов аренды"""

    class Meta:
        model = RentalPeriod
        fields = (
            'id', 'contract__number', 'start_date',
            'end_date', 'actual_price', 'duration_days'
        )
        export_order = fields


class RentalPeriodFilter(FilterSet):
    """Фильтры для периодов аренды"""

    class Meta:
        model = RentalPeriod
        fields = {
            'start_date': ['gte', 'lte'],
            'end_date': ['gte', 'lte'],
            'contract__client': ['exact'],
            'contract__asset__asset_type': ['exact']
        }


class RentalPeriodTable(tables.Table):
    """Таблица для отображения периодов аренды"""
    actions = tables.TemplateColumn(
        template_name='analytics/actions_column.html',
        orderable=False
    )

    class Meta:
        model = RentalPeriod
        template_name = "django_tables2/bootstrap5.html"
        fields = (
            'contract', 'start_date', 'end_date',
            'duration_days', 'actual_price'
        )


def update_analytics():
    """Обновление аналитических данных"""
    from .tasks import (
        update_financial_indicators,
        generate_reports,
        update_dashboards
    )

    update_financial_indicators.delay()
    generate_reports.delay()
    update_dashboards.delay()


def get_monthly_revenue(year=None):
    """Получение месячной выручки"""
    if not year:
        year = timezone.now().year

    return RentalPeriod.objects.filter(
        start_date__year=year
    ).annotate(
        month=ExtractMonth('start_date')
    ).values(
        'month'
    ).annotate(
        revenue=Sum('actual_price')
    ).order_by('month')